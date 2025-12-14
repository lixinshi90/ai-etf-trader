# -*- coding: utf-8 -*-
from __future__ import annotations
"""
导出绩效报告（Markdown + 曲线图 + Word）
- 读取项目根目录 perf_start.json 作为起始日期（默认今天）
- 从 data/trade_history.db 读取交易，筛选起始日及之后数据
- 计算关键指标（总交易、胜率、收益率、年化、最大回撤）
- 生成权益曲线图 reports/equity_curve.png
- 生成 Markdown 报告 reports/monthly_report_YYYYMMDD.md
- 生成 Word 报告 reports/monthly_report_YYYYMMDD.docx

用法：
    conda run -n ai-etf-trader python -m src.export_report
可选：设置 PERF_START（YYYY-MM-DD）环境变量覆盖 perf_start.json
"""
import json
import os
from datetime import datetime
from typing import Dict, Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches


def _project_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def _ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)


def _db_path() -> str:
    return os.path.join(_project_root(), "data", "trade_history.db")


def _read_perf_start() -> str:
    # 优先 PERF_START 环境变量
    env = os.getenv("PERF_START")
    if env:
        return env
    # 否则 perf_start.json
    fp = os.path.join(_project_root(), "perf_start.json")
    if os.path.exists(fp):
        with open(fp, "r", encoding="utf-8") as f:
            obj = json.load(f)
        return obj.get("start_date", datetime.now().strftime("%Y-%m-%d"))
    return datetime.now().strftime("%Y-%m-%d")


def _calc_metrics(trades: pd.DataFrame) -> Dict[str, Any]:
    if trades.empty:
        return {
            "total_trades": 0,
            "win_rate": 0.0,
            "total_return": 0.0,
            "annualized_return": 0.0,
            "max_drawdown": 0.0,
        }

    trades = trades.copy()
    trades["date"] = pd.to_datetime(trades["date"])  # 包含时分秒
    # 卖出为正、买入为负
    trades["profit"] = trades.apply(lambda r: r["value"] if r["action"] == "sell" else -r["value"], axis=1)

    total_trades = int((trades["action"] == "sell").sum())
    win_trades = int(((trades["action"] == "sell") & (trades["profit"] > 0)).sum())
    win_rate = (win_trades / total_trades) if total_trades > 0 else 0.0

    # 账户权益曲线（按交易后资金）
    equity = trades["capital_after"].astype(float).values
    initial_capital = float(equity[0]) if len(equity) else 100000.0
    final_capital = float(equity[-1]) if len(equity) else initial_capital
    total_return = (final_capital - initial_capital) / (initial_capital if initial_capital else 1.0)

    days = max((trades["date"].iloc[-1] - trades["date"].iloc[0]).days, 1) if not trades.empty else 1
    annualized_return = total_return * (365.0 / days) if days > 0 else 0.0

    peak = np.maximum.accumulate(equity) if len(equity) else np.array([initial_capital])
    drawdown = (equity - peak) / peak if len(equity) else np.array([0.0])
    max_drawdown = float(drawdown.min()) if len(drawdown) else 0.0

    return {
        "total_trades": total_trades,
        "win_rate": win_trades / total_trades * 100.0 if total_trades > 0 else 0.0,
        "total_return": total_return * 100.0,
        "annualized_return": annualized_return * 100.0,
        "max_drawdown": max_drawdown * 100.0,
    }


def _plot_equity(daily_dates: list[str], daily_values: list[float], out_png: str):
    _ensure_dir(os.path.dirname(out_png))
    plt.figure(figsize=(10, 4))
    plt.plot(daily_dates, daily_values, color="#2196F3", linewidth=2)
    plt.title("账户总资产曲线")
    plt.xlabel("日期")
    plt.ylabel("资产(元)")
    plt.grid(True, linestyle="--", alpha=0.3)
    plt.xticks(rotation=30)
    plt.tight_layout()
    plt.savefig(out_png, dpi=150)
    plt.close()


def _write_markdown(md_path: str, start_date: str, metrics: Dict[str, Any]):
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(f"# 模拟绩效报告（起始日：{start_date}，生成日：{datetime.now().strftime('%Y-%m-%d')}）\n\n")
        f.write("## 关键指标\n\n")
        f.write(f"- 总交易次数：{metrics['total_trades']}\n")
        f.write(f"- 胜率：{metrics['win_rate']:.2f}%\n")
        f.write(f"- 总收益率：{metrics['total_return']:.2f}%\n")
        f.write(f"- 年化收益率：{metrics['annualized_return']:.2f}%\n")
        f.write(f"- 最大回撤：{metrics['max_drawdown']:.2f}%\n\n")
        f.write("## 账户总资产曲线\n\n")
        f.write("![equity](equity_curve.png)\n\n")
        f.write("> 注：统计区间内每日最后一笔交易后的账户资金作为当日权益值。\n")


def _write_docx(docx_path: str, png_path: str, start_date: str, metrics: Dict[str, Any]):
    doc = Document()
    doc.add_heading("模拟绩效报告", level=1)
    doc.add_paragraph(f"起始日：{start_date}")
    doc.add_paragraph(f"生成日：{datetime.now().strftime('%Y-%m-%d')}")
    doc.add_heading("关键指标", level=2)
    doc.add_paragraph(f"总交易次数：{metrics['total_trades']}")
    doc.add_paragraph(f"胜率：{metrics['win_rate']:.2f}%")
    doc.add_paragraph(f"总收益率：{metrics['total_return']:.2f}%")
    doc.add_paragraph(f"年化收益率：{metrics['annualized_return']:.2f}%")
    doc.add_paragraph(f"最大回撤：{metrics['max_drawdown']:.2f}%")
    if os.path.exists(png_path):
        doc.add_heading("账户总资产曲线", level=2)
        doc.add_picture(png_path, width=Inches(6.0))
    doc.add_paragraph("注：统计区间内每日最后一笔交易后的账户资金作为当日权益值。")
    doc.save(docx_path)


def main():
    start_date = _read_perf_start()
    db_file = _db_path()
    if not os.path.exists(db_file):
        print("尚无交易数据库，无法导出报告。请先运行主任务生成交易记录。")
        return

    import sqlite3
    conn = sqlite3.connect(db_file)
    try:
        trades = pd.read_sql_query("SELECT * FROM trades ORDER BY datetime(date)", conn)
    finally:
        conn.close()

    if trades.empty:
        print("交易记录为空，暂无可导出内容。")
        return

    # 过滤起始日
    trades["date"] = pd.to_datetime(trades["date"])
    try:
        start_dt = pd.to_datetime(start_date)
    except Exception:
        start_dt = trades["date"].iloc[0]
    trades = trades[trades["date"] >= start_dt].reset_index(drop=True)

    # 计算指标
    metrics = _calc_metrics(trades)

    # 按日权益
    daily = trades.groupby(trades["date"].dt.date)["capital_after"].last().reset_index()
    daily["date"] = daily["date"].astype(str)
    dates = daily["date"].tolist()
    values = [float(v) for v in daily["capital_after"].tolist()]

    # 生成图与报告
    reports_dir = os.path.join(_project_root(), "reports")
    _ensure_dir(reports_dir)
    today = datetime.now().strftime("%Y%m%d")
    png_path = os.path.join(reports_dir, "equity_curve.png")
    _plot_equity(dates, values, png_path)

    md_path = os.path.join(reports_dir, f"monthly_report_{today}.md")
    _write_markdown(md_path, start_date, metrics)

    docx_path = os.path.join(reports_dir, f"monthly_report_{today}.docx")
    _write_docx(docx_path, png_path, start_date, metrics)

    print("已生成：", md_path)
    print("已生成：", docx_path)
    print("曲线图：", png_path)


if __name__ == "__main__":
    main()
