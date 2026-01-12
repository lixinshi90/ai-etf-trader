# -*- coding: utf-8 -*-
"""
期末综合实验报告导出（Word）
- 起始统计日：优先 PERF_START 环境变量，其次 perf_start.json 的 start_date
- 从 data/trade_history.db 读取交易，计算关键指标与生成权益曲线
- 输出：reports/期末综合实验报告.docx

用法：
    conda run -n ai-etf-trader python -m src.export_final_report
或：
    PERF_START=YYYY-MM-DD python -m src.export_final_report
"""
from __future__ import annotations

import json
import os
from datetime import datetime
from typing import Dict, Any

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches


def _project_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def _ensure_dir(p: str):
    os.makedirs(p, exist_ok=True)


def _db_path() -> str:
    return os.path.join(_project_root(), "data", "trade_history.db")


def _read_perf_start() -> str:
    env = os.getenv("PERF_START")
    if env:
        return env
    fp = os.path.join(_project_root(), "perf_start.json")
    if os.path.exists(fp):
        with open(fp, "r", encoding="utf-8") as f:
            obj = json.load(f)
        return obj.get("start_date", datetime.now().strftime("%Y-%m-%d"))
    return datetime.now().strftime("%Y-%m-%d")


def _calc_metrics(trades: pd.DataFrame) -> Dict[str, Any]:
    if trades.empty:
        return {
            "total_trades": 0,
            "win_rate": 0.0,
            "total_return": 0.0,
            "annualized_return": 0.0,
            "max_drawdown": 0.0,
        }
    trades = trades.copy()
    trades["date"] = pd.to_datetime(trades["date"])  # 含时分秒
    trades["profit"] = trades.apply(lambda r: r["value"] if r["action"] == "sell" else -r["value"], axis=1)
    total_trades = int((trades["action"] == "sell").sum())
    win_trades = int(((trades["action"] == "sell") & (trades["profit"] > 0)).sum())
    win_rate = (win_trades / total_trades * 100.0) if total_trades > 0 else 0.0

    equity = trades["capital_after"].astype(float).values
    initial_capital = float(equity[0]) if len(equity) else 100000.0
    final_capital = float(equity[-1]) if len(equity) else initial_capital
    total_return = (final_capital - initial_capital) / (initial_capital if initial_capital else 1.0) * 100.0

    days = max((trades["date"].iloc[-1] - trades["date"].iloc[0]).days, 1) if not trades.empty else 1
    annualized_return = (total_return * (365.0 / days)) if days > 0 else 0.0

    peak = np.maximum.accumulate(equity) if len(equity) else np.array([initial_capital])
    drawdown = (equity - peak) / peak if len(equity) else np.array([0.0])
    max_drawdown = float(drawdown.min()) * 100.0 if len(drawdown) else 0.0

    return {
        "total_trades": total_trades,
        "win_rate": win_rate,
        "total_return": total_return,
        "annualized_return": annualized_return,
        "max_drawdown": max_drawdown,
    }


def _plot_equity(daily_dates: list[str], daily_values: list[float], out_png: str):
    _ensure_dir(os.path.dirname(out_png))
    plt.figure(figsize=(10, 4))
    plt.plot(daily_dates, daily_values, color="#2196F3", linewidth=2)
    plt.title("账户总资产曲线")
    plt.xlabel("日期")
    plt.ylabel("资产(元)")
    plt.grid(True, linestyle="--", alpha=0.3)
    plt.xticks(rotation=30)
    plt.tight_layout()
    plt.savefig(out_png, dpi=150)
    plt.close()


def _write_docx(docx_path: str, png_path: str, start_date: str, metrics: Dict[str, Any]):
    doc = Document()
    doc.add_heading("期末综合实验报告（AI炒ETF系统）", level=1)
    doc.add_paragraph(f"起始统计日：{start_date}")
    doc.add_paragraph(f"报告生成日：{datetime.now().strftime('%Y-%m-%d')}")

    doc.add_heading("一、系统概述", level=2)
    doc.add_paragraph("本系统实现 AI 驱动的ETF量化交易：数据获取 → AI决策 → 交易执行 → 绩效评估 → Web展示 → 定时化调度。")

    doc.add_heading("二、关键绩效指标（统计区间内）", level=2)
    doc.add_paragraph(f"总交易次数：{metrics['total_trades']}")
    doc.add_paragraph(f"胜率：{metrics['win_rate']:.2f}%")
    doc.add_paragraph(f"总收益率：{metrics['total_return']:.2f}%")
    doc.add_paragraph(f"年化收益率：{metrics['annualized_return']:.2f}%")
    doc.add_paragraph(f"最大回撤：{metrics['max_drawdown']:.2f}%")

    if os.path.exists(png_path):
        doc.add_heading("三、账户总资产曲线", level=2)
        doc.add_picture(png_path, width=Inches(6.0))
        doc.add_paragraph("注：按每日最后一笔交易后的资金作为当日权益值。")

    doc.add_heading("四、实现要点", level=2)
    doc.add_paragraph("1. 数据：akshare 获取ETF日线，落地SQLite（data/etf_data.db）")
    doc.add_paragraph("2. AI 决策：OpenAI兼容接口，超时/重试、回退模型、当日缓存；决策与Prompt落盘审计")
    doc.add_paragraph("3. 交易：动态仓位与跟踪止损（开关位），成本与滑点建模；交易落地SQLite（data/trade_history.db）")
    doc.add_paragraph("4. Web：Flask API + ECharts 前端；健康检查路由 /health；/api/* 提供数据")
    doc.add_paragraph("5. 调度：main（schedule）与 daily_once（一次性）；支持Windows计划任务与VPS部署（gunicorn+nginx+systemd）")

    doc.add_heading("五、部署与运行", level=2)
    doc.add_paragraph("1. 本地：conda环境；python -m src.main / python -m src.web_app")
    doc.add_paragraph("2. 服务器：gunicorn 加载 src.web_app:app，nginx 反向代理；systemd 常驻服务；/health 健康检查")

    doc.add_heading("六、结论与展望", level=2)
    doc.add_paragraph("在保证风控的前提下，系统可稳定执行策略与可视化；后续可拓展多模型合议、备用数据源、CTA 策略等。")

    doc.save(docx_path)


def main():
    start_date = _read_perf_start()
    db_file = _db_path()
    if not os.path.exists(db_file):
        print("尚无交易数据库，无法导出报告。请先运行主任务或写入演示交易。")
        return

    import sqlite3
    conn = sqlite3.connect(db_file)
    try:
        trades = pd.read_sql_query("SELECT * FROM trades ORDER BY datetime(date)", conn)
    finally:
        conn.close()

    if trades.empty:
        print("交易记录为空，将生成零指标的空白报告。可先运行: python -m src.init_demo_trades 以生成示例交易后再导出。")
        # 仍然生成空白报告
        start_date = _read_perf_start()
        reports_dir = os.path.join(_project_root(), "reports")
        _ensure_dir(reports_dir)
        docx_path = os.path.join(reports_dir, "期末综合实验报告.docx")
        # 空指标
        metrics = {
            "total_trades": 0,
            "win_rate": 0.0,
            "total_return": 0.0,
            "annualized_return": 0.0,
            "max_drawdown": 0.0,
        }
        # 无图可画，直接写文档
        _write_docx(docx_path, png_path=os.path.join(reports_dir, "equity_curve.png"), start_date=start_date, metrics=metrics)
        print("已生成：", docx_path)
        return

    # 过滤起始日
    trades["date"] = pd.to_datetime(trades["date"])
    try:
        start_dt = pd.to_datetime(start_date)
    except Exception:
        start_dt = trades["date"].iloc[0]
    trades = trades[trades["date"] >= start_dt].reset_index(drop=True)

    # 指标
    metrics = _calc_metrics(trades)

    # 按日权益
    daily = trades.groupby(trades["date"].dt.date)["capital_after"].last().reset_index()
    daily["date"] = daily["date"].astype(str)
    dates = daily["date"].tolist()
    values = [float(v) for v in daily["capital_after"].tolist()]

    # 输出
    reports_dir = os.path.join(_project_root(), "reports")
    _ensure_dir(reports_dir)
    png_path = os.path.join(reports_dir, "equity_curve.png")
    _plot_equity(dates, values, png_path)

    docx_path = os.path.join(reports_dir, "期末综合实验报告.docx")
    _write_docx(docx_path, png_path, start_date, metrics)

    print("已生成：", docx_path)
    print("曲线图：", png_path)


if __name__ == "__main__":
    main()

