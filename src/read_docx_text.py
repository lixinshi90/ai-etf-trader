# -*- coding: utf-8 -*-
import sys
from docx import Document

def main():
    path = sys.argv[1] if len(sys.argv) > 1 else '期末综合实验报告.docx'
    doc = Document(path)
    # 打印段落
    for p in doc.paragraphs:
        t = (p.text or '').strip()
        if t:
            print(t)
    # 打印表格（如有）
    for ti, table in enumerate(doc.tables):
        print(f"\n[表格 {ti+1}]")
        for row in table.rows:
            cells = [ (c.text or '').strip() for c in row.cells ]
            print('\t'.join(cells))

if __name__ == '__main__':
    main()

